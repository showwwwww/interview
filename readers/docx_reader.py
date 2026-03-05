from __future__ import annotations

from pathlib import Path

from utils.logger import get_logger
from .base import BaseReader, Section

logger = get_logger(__name__)

_HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title"}


class DocxReader(BaseReader):
    def read(self, filepath: str) -> list[Section]:
        from docx import Document

        doc = Document(filepath)
        sections: list[Section] = []
        current_title = ""
        current_paragraphs: list[str] = []
        section_counter = 0

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            is_heading = style_name in _HEADING_STYLES

            if is_heading:
                if current_paragraphs or current_title:
                    section_counter += 1
                    sections.append(Section(
                        title=current_title or f"Section {section_counter}",
                        number=str(section_counter),
                        content="\n".join(current_paragraphs).strip(),
                    ))
                current_title = para.text.strip()
                current_paragraphs = []
            else:
                text = para.text.strip()
                if text:
                    current_paragraphs.append(text)

        if current_paragraphs or current_title:
            section_counter += 1
            sections.append(Section(
                title=current_title or f"Section {section_counter}",
                number=str(section_counter),
                content="\n".join(current_paragraphs).strip(),
            ))

        if not sections:
            full_text = "\n".join(p.text for p in doc.paragraphs).strip()
            sections.append(Section(
                title=Path(filepath).stem,
                number="1",
                content=full_text,
            ))

        logger.info("DocxReader: extracted %d section(s) from %s", len(sections), filepath)
        return sections
