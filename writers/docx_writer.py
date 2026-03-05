"""Write sections to a formatted DOCX file."""

from __future__ import annotations

from docx import Document
from docx.shared import Pt

from readers.base import Section
from utils.logger import get_logger

logger = get_logger(__name__)

_FONT_NAME = "Calibri"
_FONT_SIZE = Pt(11)


def write_docx(sections: list[Section], output_path: str) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = _FONT_NAME
    font.size = _FONT_SIZE

    for section in sections:
        heading_level = 1 if "Part" not in section.title else 2
        doc.add_heading(f"Section {section.number}: {section.title}", level=heading_level)

        if _looks_like_bullets(section.content):
            for line in section.content.split("\n"):
                line = line.strip()
                if not line:
                    continue
                clean = line.lstrip("-•*").strip()
                doc.add_paragraph(clean, style="List Bullet")
        else:
            for para_text in section.content.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    doc.add_paragraph(para_text)

    doc.save(output_path)
    logger.info("DOCX written: %s", output_path)


def _looks_like_bullets(text: str) -> bool:
    lines = [l.strip() for l in text.strip().split("\n") if l.strip()]
    if not lines:
        return False
    bullet_count = sum(1 for l in lines if l.startswith(("-", "•", "*")))
    return bullet_count / len(lines) > 0.5
