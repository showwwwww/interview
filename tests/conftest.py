"""Shared fixtures for the test suite."""

import os
import sys
from pathlib import Path
from unittest.mock import MagicMock

import pytest

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

os.environ.setdefault("OPENAI_API_KEY", "test-key-not-real")


@pytest.fixture
def tmp_output(tmp_path):
    """Return a temporary directory for output files."""
    return tmp_path / "output"


@pytest.fixture
def sample_txt_file(tmp_path):
    """Create a minimal .txt legal document."""
    content = (
        "ARTICLE I\n"
        "This agreement establishes the terms and conditions.\n\n"
        "ARTICLE II\n"
        "The parties agree to the following obligations.\n\n"
        "SECTION 3\n"
        "Termination clauses apply as described herein.\n"
    )
    p = tmp_path / "sample_legal.txt"
    p.write_text(content, encoding="utf-8")
    return p


@pytest.fixture
def sample_docx_file(tmp_path):
    """Create a minimal .docx file with headings."""
    from docx import Document

    doc = Document()
    doc.add_heading("Agreement Overview", level=1)
    doc.add_paragraph("This document outlines the terms of the agreement.")
    doc.add_heading("Obligations", level=2)
    doc.add_paragraph("Each party shall fulfill the following obligations.")
    p = tmp_path / "sample_legal.docx"
    doc.save(str(p))
    return p


@pytest.fixture
def mock_llm_client():
    """Return a mock LLMClient that echoes back simplified text."""
    client = MagicMock()
    client.convert.return_value = "This is the LLM-converted text for testing purposes."
    return client
